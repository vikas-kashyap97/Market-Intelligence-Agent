import os
import logging
from typing import List, Dict, Any
from datetime import datetime
import asyncio
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from config.settings import Settings

logger = logging.getLogger(__name__)

class ReportExporter:
    """Handle export of reports to various formats"""
    
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
        # Ensure exports directory exists
        self.exports_dir = os.path.join(output_dir, "exports")
        os.makedirs(self.exports_dir, exist_ok=True)
    
    async def export_to_pdf(self, content: str, chart_files: List[str], title: str) -> str:
        """Export report to PDF format"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            filename = f"report_{timestamp}.pdf"
            filepath = os.path.join(self.exports_dir, filename)
            
            # Create PDF document
            doc = SimpleDocTemplate(filepath, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=HexColor('#2E86AB'),
                alignment=1  # Center alignment
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceBefore=20,
                spaceAfter=12,
                textColor=HexColor('#A23B72')
            )
            
            # Add title
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))
            
            # Process markdown content
            lines = content.split('\n')
            current_section = []
            
            for line in lines:
                if line.startswith('# '):
                    # Main title - skip as we already added it
                    continue
                elif line.startswith('## '):
                    # Section heading
                    if current_section:
                        story.extend(self._process_section(current_section, styles))
                        current_section = []
                    
                    heading_text = line[3:].strip()
                    story.append(Paragraph(heading_text, heading_style))
                    story.append(Spacer(1, 12))
                elif line.startswith('### '):
                    # Subsection
                    subsection_text = line[4:].strip()
                    story.append(Paragraph(subsection_text, styles['Heading3']))
                    story.append(Spacer(1, 8))
                elif line.strip():
                    current_section.append(line)
            
            # Process remaining section
            if current_section:
                story.extend(self._process_section(current_section, styles))
            
            # Add charts
            if chart_files:
                story.append(PageBreak())
                story.append(Paragraph("Charts and Visualizations", heading_style))
                story.append(Spacer(1, 20))
                
                for chart_file in chart_files:
                    if chart_file.endswith('.png'):
                        chart_path = os.path.join(self.output_dir, chart_file)
                        if os.path.exists(chart_path):
                            try:
                                # Add chart with caption
                                story.append(Paragraph(f"Chart: {chart_file.replace('_', ' ').title()}", styles['Heading4']))
                                story.append(Spacer(1, 8))
                                
                                img = Image(chart_path, width=6*inch, height=4*inch)
                                story.append(img)
                                story.append(Spacer(1, 20))
                            except Exception as e:
                                logger.warning(f"Could not add chart {chart_file} to PDF: {str(e)}")
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"Successfully exported PDF: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Failed to export PDF: {str(e)}")
            raise
    
    async def export_to_docx(self, content: str, chart_files: List[str], title: str) -> str:
        """Export report to DOCX format"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            filename = f"report_{timestamp}.docx"
            filepath = os.path.join(self.exports_dir, filename)
            
            # Create document
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Process markdown content
            lines = content.split('\n')
            current_paragraph = []
            
            for line in lines:
                if line.startswith('# '):
                    # Main title - skip as we already added it
                    continue
                elif line.startswith('## '):
                    # Section heading
                    if current_paragraph:
                        doc.add_paragraph('\n'.join(current_paragraph))
                        current_paragraph = []
                    
                    heading_text = line[3:].strip()
                    doc.add_heading(heading_text, level=1)
                elif line.startswith('### '):
                    # Subsection
                    if current_paragraph:
                        doc.add_paragraph('\n'.join(current_paragraph))
                        current_paragraph = []
                    
                    subsection_text = line[4:].strip()
                    doc.add_heading(subsection_text, level=2)
                elif line.startswith('**') and line.endswith('**'):
                    # Bold text
                    if current_paragraph:
                        doc.add_paragraph('\n'.join(current_paragraph))
                        current_paragraph = []
                    
                    bold_text = line[2:-2]
                    p = doc.add_paragraph()
                    p.add_run(bold_text).bold = True
                elif line.startswith('- '):
                    # Bullet point
                    if current_paragraph:
                        doc.add_paragraph('\n'.join(current_paragraph))
                        current_paragraph = []
                    
                    bullet_text = line[2:].strip()
                    doc.add_paragraph(bullet_text, style='List Bullet')
                elif line.strip():
                    current_paragraph.append(line)
                else:
                    # Empty line
                    if current_paragraph:
                        doc.add_paragraph('\n'.join(current_paragraph))
                        current_paragraph = []
            
            # Add remaining paragraph
            if current_paragraph:
                doc.add_paragraph('\n'.join(current_paragraph))
            
            # Add charts
            if chart_files:
                doc.add_page_break()
                doc.add_heading('Charts and Visualizations', level=1)
                
                for chart_file in chart_files:
                    if chart_file.endswith('.png'):
                        chart_path = os.path.join(self.output_dir, chart_file)
                        if os.path.exists(chart_path):
                            try:
                                # Add chart caption
                                doc.add_heading(f"Chart: {chart_file.replace('_', ' ').title()}", level=2)
                                
                                # Add image
                                doc.add_picture(chart_path, width=Inches(6))
                                doc.add_paragraph()  # Add space after image
                            except Exception as e:
                                logger.warning(f"Could not add chart {chart_file} to DOCX: {str(e)}")
            
            # Save document
            doc.save(filepath)
            
            logger.info(f"Successfully exported DOCX: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Failed to export DOCX: {str(e)}")
            raise
    
    def _process_section(self, lines: List[str], styles) -> List:
        """Process a section of content for PDF"""
        elements = []
        
        for line in lines:
            if line.startswith('**') and line.endswith('**'):
                # Bold text
                bold_text = line[2:-2]
                elements.append(Paragraph(f"<b>{bold_text}</b>", styles['Normal']))
            elif line.startswith('- '):
                # Bullet point
                bullet_text = line[2:].strip()
                elements.append(Paragraph(f"• {bullet_text}", styles['Normal']))
            elif line.strip():
                # Regular paragraph
                elements.append(Paragraph(line, styles['Normal']))
            
            elements.append(Spacer(1, 6))
        
        return elements
    
    async def export_to_notion(self, content: str, title: str) -> str:
        """Export report to Notion (if API key is available)"""
        try:
            if not Settings.NOTION_API_KEY:
                logger.warning("Notion API key not available, skipping Notion export")
                return "Notion export skipped - API key not configured"
            
            # This would implement Notion API integration
            # For now, return a placeholder
            logger.info("Notion export would be implemented here")
            return "Notion export feature coming soon"
            
        except Exception as e:
            logger.error(f"Failed to export to Notion: {str(e)}")
            return f"Notion export failed: {str(e)}"
